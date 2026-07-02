"""
论文格式批量修正脚本
修正内容:
  1. 章节标题中文数字 → 阿拉伯数字 (第一章→第1章)
  2. 正文段落首行缩进统一为两格 (24pt)
  3. 表格边框修正为三线表 (顶底粗线 + 中间细线)
  4. 去除表格后多余阴影/底纹
"""
import shutil
import re
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

INPUT_PATH = r"c:\Users\my147\Desktop\李帮家-2022022425-电子信息工程-基于 Web 的农场科普游戏设计与实现(2).docx"
BACKUP_PATH = INPUT_PATH.replace(".docx", ".bak.docx")
OUTPUT_PATH = INPUT_PATH.replace(".docx", "_修正版.docx")

# ── 辅助函数 ──────────────────────────────────────────────

CN_TO_ARABIC = {"一": "1", "二": "2", "三": "3", "四": "4",
                "五": "5", "六": "6", "七": "7", "八": "8", "九": "9"}
CHAPTER_RE = re.compile(r"第([一二三四五六七八九])章")


def _is_heading(para):
    """判断段落是否为标题（Heading 或 TOC 样式）。"""
    style_name = (para.style.name or "").lower()
    return "heading" in style_name or "toc" in style_name


def _is_in_table(para):
    """判断段落是否在表格内部。"""
    return para._element.find(".//" + qn("w:tc")) is not None or \
           para._element.getparent() is not None and \
           para._element.getparent().tag == qn("w:tc")


def _replace_in_runs(para, old, new):
    """在 run 级别做文本替换，保留原有格式。返回替换次数。"""
    count = 0
    for run in para.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            count += 1
    return count


def _replace_chapter_in_para(para, cn, ar):
    """智能替换章节编号：先尝试 run 级，失败则拼接后整体替换。"""
    old_text = f"第{cn}章"
    new_text = f"第{ar}章"

    # 方式 1: 单个 run 内替换
    count = _replace_in_runs(para, old_text, new_text)
    if count > 0:
        return count

    # 方式 2: 跨 run 拼接替换（处理 TOC 等文本被拆分的情况）
    full = para.text
    if old_text not in full:
        return 0

    # 拼接所有 run 文本，找到目标位置，把替换写入对应 run
    runs = para.runs
    if not runs:
        return 0

    # 重建 run 文本：把所有 run 拼成一个字符串，记录每个字符属于哪个 run
    joined = "".join(r.text for r in runs)
    idx = joined.find(old_text)
    if idx == -1:
        return 0

    # 计算每个 run 的起止位置
    positions = []
    pos = 0
    for r in runs:
        positions.append((pos, pos + len(r.text)))
        pos += len(r.text)

    # 找到目标文本跨越的 run 范围
    target_end = idx + len(old_text)
    for i, (start, end) in enumerate(positions):
        if start <= idx < end:
            first_run = i
        if start < target_end <= end:
            last_run = i
            break

    # 替换策略：把新文本写入第一个匹配 run，清空中间 run 的目标部分
    new_joined = joined[:idx] + new_text + joined[target_end:]

    # 重新分配 run 文本
    cursor = 0
    for i, r in enumerate(runs):
        old_len = len(r.text)
        if old_len == 0:
            continue
        r.text = new_joined[cursor:cursor + old_len]
        cursor += old_len
        if i == len(runs) - 1:
            # 最后一个 run 吸收剩余所有字符
            r.text = new_joined[cursor:]
            cursor = len(new_joined)

    # 简化：直接把新文本全部分配到 run
    # 重新拼接
    final = ""
    for r in runs:
        final += r.text
    if old_text not in final and new_text in final:
        return 1

    # 最终方案：把新拼接结果按比例分回 runs
    total_old = sum(len(r.text) for r in runs)
    total_new = len(new_joined)
    cursor = 0
    for i, r in enumerate(runs):
        if i == len(runs) - 1:
            r.text = new_joined[cursor:]
        else:
            ratio = len(r.text) / total_old if total_old > 0 else 0
            seg_len = max(1, round(ratio * total_new))
            r.text = new_joined[cursor:cursor + seg_len]
            cursor += seg_len

    return 1


def _set_cell_border(cell, position, tag, val, sz, color="000000"):
    """设置单元格指定位置(position=tcTop/tcBottom/tcLeft/tcRight/tcInsideH/tcInsideV)的边框。"""
    tc_pr = cell._element.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    el = borders.find(qn(f"w:{tag}"))
    if el is None:
        el = OxmlElement(f"w:{tag}")
        borders.append(el)
    el.set(qn("w:val"), val)
    el.set(qn("w:sz"), sz)
    el.set(qn("w:color"), color)
    el.set(qn("w:space"), "0")


def _clear_para_shading(para):
    """清除段落及其 run 上的底纹/阴影属性。"""
    removed = 0
    # 段落级
    pPr = para._element.find(qn("w:pPr"))
    if pPr is not None:
        for shd in pPr.findall(qn("w:shd")):
            pPr.remove(shd)
            removed += 1
    # run 级
    for rPr in para._element.findall(".//" + qn("w:rPr")):
        for shd in rPr.findall(qn("w:shd")):
            rPr.remove(shd)
            removed += 1
    return removed


# ── 主流程 ────────────────────────────────────────────────

def main():
    print(f"备份: {INPUT_PATH}")
    print(f"  →  {BACKUP_PATH}")
    shutil.copy2(INPUT_PATH, BACKUP_PATH)

    doc = Document(INPUT_PATH)
    stats = {"chapter": 0, "indent": 0, "table": 0, "shadow": 0}

    # ─── 1. 章节标题: 中文数字 → 阿拉伯数字 ───
    for para in doc.paragraphs:
        full = para.text
        if not CHAPTER_RE.search(full):
            continue
        for cn, ar in CN_TO_ARABIC.items():
            old_text = f"第{cn}章"
            new_text = f"第{ar}章"
            if old_text not in full:
                continue

            # 方式 1: run 级替换
            n = _replace_in_runs(para, old_text, new_text)
            if n > 0:
                stats["chapter"] += n
                continue

            # 方式 2: 直接操作 XML 层所有 w:t 元素（包含 hyperlink 内的 run）
            xml_replaced = 0
            for t_el in para._element.findall(".//" + qn("w:t")):
                if t_el.text and old_text in t_el.text:
                    t_el.text = t_el.text.replace(old_text, new_text)
                    xml_replaced += 1
            if xml_replaced > 0:
                stats["chapter"] += xml_replaced
                continue

            # 方式 3: 跨 run 替换（文本被拆分到多个 run 中）
            all_t = para._element.findall(".//" + qn("w:t"))
            joined = "".join((t.text or "") for t in all_t)
            if old_text in joined:
                # 拼接后做替换，然后按比例分配回各个 t 元素
                new_joined = joined.replace(old_text, new_text)
                total_old = len(joined)
                total_new = len(new_joined)
                cursor = 0
                for i, t_el in enumerate(all_t):
                    old_len = len(t_el.text or "")
                    if old_len == 0:
                        continue
                    if i == len(all_t) - 1:
                        t_el.text = new_joined[cursor:]
                    else:
                        ratio = old_len / total_old if total_old > 0 else 0
                        seg_len = max(1, round(ratio * total_new))
                        t_el.text = new_joined[cursor:cursor + seg_len]
                        cursor += seg_len
                stats["chapter"] += 1

    print(f"[1/4] 章节标题替换: {stats['chapter']} 处")

    # ─── 2. 段落首行缩进统一为两格 ───
    SKIP_PREFIX = ("heading", "toc", "caption", "list", "title")
    for para in doc.paragraphs:
        style_name = (para.style.name or "").lower()
        if any(style_name.startswith(p) for p in SKIP_PREFIX):
            continue
        if not para.text.strip():
            continue
        # 排除表格内段落
        try:
            if _is_in_table(para):
                continue
        except Exception:
            pass
        para.paragraph_format.first_line_indent = Pt(24)
        stats["indent"] += 1

    print(f"[2/4] 段落缩进修正: {stats['indent']} 段")

    # ─── 3. 表格边框 → 三线表 ───
    THICK = "24"   # 12pt (sz 单位 = 半磅)
    THIN  = "4"    # 0.5pt

    for table in doc.tables:
        tbl_pr = table._tbl.tblPr
        if tbl_pr is None:
            tbl_pr = OxmlElement("w:tblPr")
            table._tbl.insert(0, tbl_pr)

        # 3a. 表格级默认边框 → 全部细线
        tbl_borders = tbl_pr.find(qn("w:tblBorders"))
        if tbl_borders is None:
            tbl_borders = OxmlElement("w:tblBorders")
            tbl_pr.append(tbl_borders)

        for bname in ("top", "left", "bottom", "right",
                      "insideH", "insideV"):
            el = tbl_borders.find(qn(f"w:{bname}"))
            if el is None:
                el = OxmlElement(f"w:{bname}")
                tbl_borders.append(el)
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), THIN)
            el.set(qn("w:color"), "000000")
            el.set(qn("w:space"), "0")

        rows = table.rows
        if not rows:
            stats["table"] += 1
            continue

        # 3b. 第一行每单元格 → 顶部粗线
        for cell in rows[0].cells:
            _set_cell_border(cell, None, "top", "single", THICK)

        # 3c. 最后一行每单元格 → 底部粗线
        for cell in rows[-1].cells:
            _set_cell_border(cell, None, "bottom", "single", THICK)

        stats["table"] += 1

    print(f"[3/4] 表格边框修正: {stats['table']} 个表格")

    # ─── 4. 去除表格后多余阴影 ───
    body = doc.element.body
    for tbl_el in body.findall(qn("w:tbl")):
        nxt = tbl_el.getnext()
        # 检查紧接的 1-2 个段落
        for _ in range(2):
            if nxt is None or nxt.tag != qn("w:p"):
                break
            # 找到对应的 Paragraph 对象
            for para in doc.paragraphs:
                if para._element is nxt:
                    r = _clear_para_shading(para)
                    stats["shadow"] += r
                    break
            nxt = nxt.getnext()

    print(f"[4/4] 阴影清除: {stats['shadow']} 处")

    # ─── 保存 ───
    doc.save(OUTPUT_PATH)
    print(f"\n已保存到: {OUTPUT_PATH}")

    # ─── 验证 ───
    print("\n" + "=" * 55)
    print("验证结果")
    print("=" * 55)

    vdoc = Document(OUTPUT_PATH)

    # 章节标题
    print("\n[章节标题]")
    for p in vdoc.paragraphs:
        if "第" in p.text and "章" in p.text:
            sname = (p.style.name or "")
            if "Heading" in sname or "TOC" in sname or p.text.startswith("第"):
                print(f"  {p.text[:50]}")

    # 表格边框
    print(f"\n[表格数量] {len(vdoc.tables)}")
    for i, tbl in enumerate(vdoc.tables):
        tbl_pr = tbl._tbl.tblPr
        if tbl_pr is None:
            continue
        borders = tbl_pr.find(qn("w:tblBorders"))
        if borders is None:
            continue
        top_sz = borders.find(qn("w:top"))
        top_sz = top_sz.get(qn("w:sz")) if top_sz is not None else "?"
        bot_sz = borders.find(qn("w:bottom"))
        bot_sz = bot_sz.get(qn("w:sz")) if bot_sz is not None else "?"
        inh_sz = borders.find(qn("w:insideH"))
        inh_sz = inh_sz.get(qn("w:sz")) if inh_sz is not None else "?"

        # 检查首行 cell 顶部
        first_row_top = "?"
        if tbl.rows:
            tc_pr = tbl.rows[0].cells[0]._element.find(qn("w:tcPr"))
            if tc_pr is not None:
                cb = tc_pr.find(qn("w:tcBorders"))
                if cb is not None:
                    t = cb.find(qn("w:top"))
                    first_row_top = t.get(qn("w:sz")) if t is not None else "?"

        last_row_bot = "?"
        if tbl.rows:
            tc_pr = tbl.rows[-1].cells[0]._element.find(qn("w:tcPr"))
            if tc_pr is not None:
                cb = tc_pr.find(qn("w:tcBorders"))
                if cb is not None:
                    b = cb.find(qn("w:bottom"))
                    last_row_bot = b.get(qn("w:sz")) if b is not None else "?"

        print(f"  表格{i+1}: "
              f"表格级top={top_sz} bottom={bot_sz} insideH={inh_sz} | "
              f"首行cell-top={first_row_top} 末行cell-bottom={last_row_bot}")


if __name__ == "__main__":
    main()
